import os
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from flask import Flask, render_template, request, send_file

# ──────────────────────────────────────────────
# Config
# ──────────────────────────────────────────────
TEMPLATE_PATH = os.getenv("TEMPLATE_PATH", "jrsu_word_file.docx")
OUTPUT_DIR    = os.getenv("OUTPUT_DIR",    "output")
UPLOAD_DIR    = os.getenv("UPLOAD_DIR",    "uploads")
# LibreOffice path — only used in Docker/Linux; on Windows we use docx2pdf instead
LIBREOFFICE   = os.getenv("LIBREOFFICE",   "soffice")
PORT          = int(os.getenv("PORT", 8080))

PLACEHOLDER_RE = re.compile(r"\{\{(.*?)\}\}")

# Create directories on startup
Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)

app = Flask(__name__, static_folder="src/main/resources/static")
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5 MB upload limit


# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def detect_placeholders() -> list:
    """Read the DOCX template and return sorted list of {{placeholder}} names."""
    found = set()
    try:
        doc = Document(TEMPLATE_PATH)

        # Body paragraphs
        for para in doc.paragraphs:
            for m in PLACEHOLDER_RE.finditer(para.text):
                name = m.group(1).strip()
                if name:
                    found.add(name)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for m in PLACEHOLDER_RE.finditer(para.text):
                            name = m.group(1).strip()
                            if name:
                                found.add(name)

        # Headers and footers
        for section in doc.sections:
            for hdr in [section.header, section.footer]:
                if hdr:
                    for para in hdr.paragraphs:
                        for m in PLACEHOLDER_RE.finditer(para.text):
                            name = m.group(1).strip()
                            if name:
                                found.add(name)

    except Exception as e:
        print(f"WARNING: Could not read template: {e}")

    return sorted(found)


def _replace_in_paragraph(para, context):
    """Replace {{placeholder}} in a paragraph handling runs split by Word."""
    full = "".join(run.text or "" for run in para.runs)
    if "{{" not in full:
        return

    replaced = full
    for key, val in context.items():
        replaced = replaced.replace("{{" + key + "}}", val)
        replaced = replaced.replace("{{ " + key + " }}", val)   # with spaces

    if replaced == full:
        return

    # Write replacement into first run, clear the rest
    if para.runs:
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ""


def fill_docx(context, logo_path):
    """Fill template placeholders and optionally swap the logo image."""
    filled_path = str(Path(OUTPUT_DIR) / "output_filled.docx")
    doc = Document(TEMPLATE_PATH)

    # Body
    for para in doc.paragraphs:
        _replace_in_paragraph(para, context)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, context)

    # Headers / footers
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    _replace_in_paragraph(para, context)

    doc.save(filled_path)

    # Swap logo via zip manipulation
    if logo_path and Path(logo_path).is_file():
        _replace_logo_in_docx(filled_path, logo_path)

    return filled_path


def _replace_logo_in_docx(docx_path, logo_path):
    """Replace word/media/image1.png inside the DOCX zip with a custom logo."""
    tmp_path = docx_path + ".tmp"
    try:
        logo_bytes = Path(logo_path).read_bytes()
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/media/image1.png":
                    zout.writestr(item, logo_bytes)
                else:
                    zout.writestr(item, data)
        shutil.move(tmp_path, docx_path)
    except Exception as e:
        print(f"WARNING: Could not replace logo: {e}")


def convert_to_pdf(docx_path):
    """
    Convert DOCX → PDF.

    Strategy (tried in order):
      1. docx2pdf  — uses MS Word COM on Windows (best quality, no extra install)
      2. LibreOffice headless — works on Linux/Docker
    """
    output_pdf = str(Path(OUTPUT_DIR) / "output_filled.pdf")

    # ── Option 1: docx2pdf (Windows / macOS with Word installed) ──────────
    try:
        from docx2pdf import convert as _d2p_convert
        _d2p_convert(str(Path(docx_path).resolve()), output_pdf)
        if Path(output_pdf).exists():
            print("PDF created via docx2pdf")
            return output_pdf
    except Exception as e:
        print(f"docx2pdf failed ({e}), trying LibreOffice …")

    # ── Option 2: LibreOffice headless ────────────────────────────────────
    abs_docx   = str(Path(docx_path).resolve())
    abs_outdir = str(Path(OUTPUT_DIR).resolve())

    # Try common Windows install paths if 'soffice' is not on PATH
    lo_candidates = [
        LIBREOFFICE,
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for lo in lo_candidates:
        if not Path(lo).exists() and lo != LIBREOFFICE:
            continue
        result = subprocess.run(
            [lo, "--headless", "--convert-to", "pdf",
             abs_docx, "--outdir", abs_outdir],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            print(f"PDF created via LibreOffice ({lo})")
            return output_pdf
        print(f"LibreOffice at '{lo}' failed: {result.stderr.strip()}")

    raise RuntimeError(
        "PDF conversion failed. Neither docx2pdf nor LibreOffice produced a PDF.\n"
        "Make sure Microsoft Word (Windows) or LibreOffice is installed."
    )


def add_border_to_pdf(pdf_path):
    """Draw a black rectangle border on every page using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        margin    = 28    # ~1 cm
        linewidth = 1.5

        tmp_path = pdf_path + ".bordered.tmp"
        doc = fitz.open(pdf_path)
        for page in doc:
            r = page.rect
            rect = fitz.Rect(
                r.x0 + margin, r.y0 + margin,
                r.x1 - margin, r.y1 - margin
            )
            page.draw_rect(rect, color=(0, 0, 0), width=linewidth)
        doc.save(tmp_path)
        doc.close()
        shutil.move(tmp_path, pdf_path)
        print("Border added to PDF")
    except ImportError:
        print("WARNING: PyMuPDF not installed; skipping border.")
    except Exception as e:
        print(f"WARNING: Could not add border: {e}")


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    placeholders = detect_placeholders()
    error = "No placeholders detected in DOCX template." if not placeholders else None
    return render_template(
        "index.html",
        placeholders=placeholders,
        error=error,
        pdf_ready=False
    )


@app.route("/generate", methods=["POST"])
def generate():
    placeholders = detect_placeholders()

    # Collect form values for each placeholder
    context = {ph: request.form.get(ph, "") for ph in placeholders}

    # Handle logo upload
    logo_path = None
    logo_file = request.files.get("logoFile")
    if logo_file and logo_file.filename:
        ext = Path(logo_file.filename).suffix or ".png"
        save_path = Path(UPLOAD_DIR) / ("custom_logo" + ext)
        try:
            logo_file.save(str(save_path))
            logo_path = str(save_path)
        except Exception as e:
            print(f"WARNING: Could not save logo: {e}")

    # Fill DOCX → convert to PDF
    try:
        filled_docx = fill_docx(context, logo_path)
        convert_to_pdf(filled_docx)

        # Optional border
        if request.form.get("addBorder", "false").lower() == "true":
            pdf_path = str(Path(OUTPUT_DIR) / "output_filled.pdf")
            add_border_to_pdf(pdf_path)

        return render_template(
            "index.html",
            placeholders=placeholders,
            pdf_ready=True,
            error=None
        )

    except Exception as e:
        return render_template(
            "index.html",
            placeholders=placeholders,
            pdf_ready=False,
            error=f"PDF generation failed: {e}"
        )


@app.route("/download")
def download():
    pdf_path = Path(OUTPUT_DIR) / "output_filled.pdf"
    if not pdf_path.exists():
        return "PDF not found", 404
    return send_file(
        str(pdf_path.resolve()),
        mimetype="application/pdf",
        as_attachment=True,
        download_name="assignment.pdf"
    )


@app.route("/view-pdf")
def view_pdf():
    pdf_path = Path(OUTPUT_DIR) / "output_filled.pdf"
    if not pdf_path.exists():
        return "PDF not found", 404
    return send_file(str(pdf_path.resolve()), mimetype="application/pdf")


# ──────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT, debug=False)
